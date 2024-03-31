import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_document(folder_path, output_file):
    doc = Document()

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            # Create a new paragraph for each file
            p = doc.add_paragraph()
            # Set paragraph alignment to left
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Add the filename in bold
            run = p.add_run(file)
            run.bold = True
            run.font.size = Pt(12)
            p.add_run("\n")  # Add a newline after the filename
            
            # Read file contents and add to the document
            with open(file_path, 'r') as f:
                doc.add_paragraph(f.read())

    # Save the document
    doc.save(output_file)

# Replace 'path_to_your_code_folder' with the path to your code folder
code_folder = 'lib'

# Replace 'output_document.docx' with the desired output file name
output_file = './output_document.docx'

generate_document(code_folder, output_file)
